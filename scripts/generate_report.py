#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Bot_IT Project Report Generator
Generates a professional Word document for the university chatbot project
"""

import sys
import io
# Set UTF-8 encoding for output
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_background(cell, color):
    """Set background color for a table cell"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_page_number(section):
    """Add page numbers to the document"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL (Right-to-Left) for Arabic text"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_indent = Inches(0)
    for run in paragraph.runs:
        run.font.rtl = True

def create_bot_it_report():
    """Create the complete Bot_IT project report"""
    
    # Create a new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "تقرير مشروع Bot_IT - روبوت الجامعة التقني"
    doc.core_properties.author = "فريق المشروع"
    doc.core_properties.subject = "Voice + Chatbot University Project"
    
    # Set up styles
    styles = doc.styles
    
    # Normal style - Arabic font
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Cairo'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    
    # Heading 1 style
    heading1 = styles['Heading 1']
    heading1.font.name = 'Cairo'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Heading 2 style
    heading2 = styles['Heading 2']
    heading2.font.name = 'Cairo'
    heading2.font.size = Pt(16)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 76, 153)  # Medium blue
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Heading 3 style
    heading3 = styles['Heading 3']
    heading3.font.name = 'Cairo'
    heading3.font.size = Pt(14)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(51, 102, 153)  # Light blue
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ============ COVER PAGE ============
    # Add cover page section
    cover = doc.add_section()
    cover.orientation = WD_ORIENTATION.PORTRAIT
    
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('اسم الجامعة: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Project title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run().add_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('تقرير مشروع')
    run.font.name = 'Cairo'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Bot_IT – روبوت الجامعة التقني')
    run.font.name = 'Cairo'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 76, 153)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(Voice + Chatbot)')
    run.font.name = 'Cairo'
    run.font.size = Pt(14)
    run.font.italic = True
    
    # Team members
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('أعضاء الفريق:')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('1. ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2. ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('3. ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('4. ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    # Supervisor
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('المشرف: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    run.font.bold = True
    
    # Date
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('التاريخ: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    
    # Add page break after cover
    add_page_break(doc)
    
    # ============ TABLE OF CONTENTS ============
    p = doc.add_paragraph('فهرس المحتويات', style='Heading 1')
    
    p = doc.add_paragraph()
    run = p.add_run('[سيتم إنشاء الفهرس تلقائياً عند تحديث المستند في Microsoft Word]')
    run.font.name = 'Cairo'
    run.font.size = Pt(11)
    run.font.italic = True
    
    # Manual table of contents for now
    toc_items = [
        ('1. ملخص تنفيذي', 2),
        ('2. مقدمة', 3),
        ('3. فكرة المشروع ومشكلة البحث', 3),
        ('4. أهداف المشروع', 4),
        ('5. نطاق المشروع', 5),
        ('6. المستخدمون المستهدفون', 6),
        ('7. المتطلبات', 6),
        ('8. التقنيات والأدوات المستخدمة', 8),
        ('9. المعمارية', 9),
        ('10. تصميم واجهة المستخدم', 10),
        ('11. الأمان والخصوصية', 11),
        ('12. خطة الاختبار', 12),
        ('13. النتائج المتوقعة ومؤشرات النجاح', 13),
        ('14. التحديات والمشاكل والحلول', 13),
        ('15. التحسينات المستقبلية', 14),
        ('16. خاتمة', 15),
        ('17. ملاحق', 15),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item + ' - ' + f'الصفحة {page}')
    
    add_page_break(doc)
    
    # ============ 1. EXECUTIVE SUMMARY ============
    doc.add_heading('1. ملخص تنفيذي', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'يقدم هذا التقرير تفاصيل مشروع Bot_IT – روبوت الجامعة التقني، وهو نظام دردشة صوتي ونصي '
        'ذكي مصمم خصيصاً للإجابة على الأسئلة التقنية والتكنولوجية داخل البيئة الجامعية. '
        'يتميز المشروع باستخدام تقنيات الذكاء الاصطناعي الحديثة، بما في ذلك نماذج Gemini/Vertex AI '
        'لتوليد الإجابات، مع واجهة تفاعلية تدعم اللغة العربية وتعمل على المتصفح مباشرة.'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'يعتمد النظام على معمارية WebSocket للاتصال الفوري بين الواجهة الأمامية والخلفية، '
        'مع استخدام Web Speech API للتعرف على الصوت وتحويل النص إلى صوت. '
        'يتميز الروبوت بميزة "كلمة التنبيه" حيث يستجيب فقط للأسئلة التي تبدأ بكلمة "روبوت"، '
        'مما يضمن تجربة مستخدم طبيعية وفعالة.'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'يهدف المشروع إلى توفير دعم تقني فوري للطلاب وأعضاء الهيئة التدريسية، '
        'مع التركيز على الأسئلة المتعلقة بالبرمجة والشبكات وقواعد البيانات والذكاء الاصطناعي '
        'وغيرها من المجالات التقنية.'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    add_page_break(doc)
    
    # ============ 2. INTRODUCTION ============
    doc.add_heading('2. مقدمة', level=1)
    
    doc.add_heading('2.1 خلفية المشروع', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'في عصر التحول الرقمي المتسارع، أصبحت الأنظمة الذكية جزءاً لا يتجزأ من الحياة الأكاديمية. '
        'يواجه الطلاب وأعضاء الهيئة التدريسية تحديات يومية تتطلب إجابات فورية ودقيقة لأسئلتهم التقنية. '
        'من هنا جاءت فكرة مشروع Bot_IT كحل مبتكر يوفر دعماً تقنياً على مدار الساعة.'
    )
    run.font.name = 'Cairo'
    
    doc.add_heading('2.2 أهمية المشروع', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'تتمثل أهمية المشروع في:'
    )
    run.font.name = 'Cairo'
    
    # Bullet points
    points = [
        'توفير الوقت والجهد على الطلاب والباحثين',
        'تقديم إجابات دقيقة وموثوقة للاستفسارات التقنية',
        'دعم اللغة العربية والتفاعل الصوتي',
        'سهولة الاستخدام والوصول عبر المتصفح',
        'قابلية التطوير والتوسع مستقبلاً'
    ]
    
    for point in points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 3. PROJECT IDEA AND RESEARCH PROBLEM ============
    doc.add_heading('3. فكرة المشروع ومشكلة البحث', level=1)
    
    doc.add_heading('3.1 المشكلة التي يحلها المشروع', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'يواجه الطلاب في الجامعات التقنية صعوبة في الحصول على إجابات فورية لأسئلتهم التقنية، '
        'خاصة خارج أوقات الدوام الرسمي. كما أن البحث في المصادر المتعددة قد يستغرق وقتاً طويلاً، '
        'وقد لا تكون الإجابات دقيقة أو محدثة.'
    )
    run.font.name = 'Cairo'
    
    doc.add_heading('3.2 لماذا روبوت تقني مفيد للطلاب؟', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'يقدم الروبوت التقني عدة فوائد:'
    )
    run.font.name = 'Cairo'
    
    benefits = [
        'إجابات فورية على مدار الساعة',
        'دقة عالية في المعلومات التقنية',
        'واجهة سهلة الاستخدام تدعم الصوت والنص',
        'توفير مصدر موثوق للمعلومات التقنية',
        'مساعدة في التعلم الذاتي والتطوير المهني'
    ]
    
    for benefit in benefits:
        p = doc.add_paragraph(benefit, style='List Bullet')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 4. PROJECT OBJECTIVES ============
    doc.add_heading('4. أهداف المشروع', level=1)
    
    doc.add_heading('4.1 الأهداف العامة', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'يهدف المشروع إلى تطوير نظام روبوت محادثة ذكي قادر على فهم والرد على '
        'الاستفسارات التقنية باللغة العربية بشكل فعال ودقيق.'
    )
    run.font.name = 'Cairo'
    
    doc.add_heading('4.2 الأهداف التفصيلية', level=2)
    
    objectives = [
        'تصميم وتطوير واجهة مستخدم عربية بسيطة وجذابة',
        'بناء خادم خلفي باستخدام Node.js وWebSocket',
        'دمج نموذج Gemini/Vertex AI لتوليد الإجابات الذكية',
        'تنفيذ التعرف على الصوت وتحويل النص إلى صوت باستخدام Web Speech API',
        'تطبيق نظام كلمة التنبيه "روبوت" لتحسين تجربة المستخدم',
        'ضمان أمان البيانات والخصوصية',
        'اختبار النظام بشكل شامل قبل الإطلاق'
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(f'{i}. {obj}')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 5. PROJECT SCOPE ============
    doc.add_heading('5. نطاق المشروع', level=1)
    
    doc.add_heading('5.1 ماذا يجيب الروبوت؟', level=2)
    p = doc.add_paragraph()
    run = p.add_run('يختص الروبوت بالإجابة على الأسئلة في المجالات التقنية التالية:')
    run.font.name = 'Cairo'
    
    # Create table for technical topics
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'المجال التقني'
    hdr_cells[1].text = 'أمثلة على الأسئلة'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Data rows
    topics = [
        ('البرمجة', 'ما هي Python؟ كيف أكتب دالة في JavaScript؟'),
        ('الشبكات', 'ما هو بروتوكول HTTP؟ كيف أعمل IP Address؟'),
        ('قواعد البيانات', 'ما الفرق بين SQL وNoSQL؟ كيف أستخدم MySQL؟'),
        ('الذكاء الاصطناعي', 'ما هو التعلم العميق؟ كيف أعمل نموذج ML؟'),
        ('أنظمة التشغيل', 'كيف أستخدم أوامر Linux؟ ما هو Kernel؟'),
        ('أمن المعلومات', 'ما هي أفضل ممارسات الأمان؟ كيف أحمي كلمات المرور؟'),
    ]
    
    for topic, examples in topics:
        row_cells = table.add_row().cells
        row_cells[0].text = topic
        row_cells[1].text = examples
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
                    run.font.size = Pt(10)
    
    doc.add_paragraph().add_run().add_break()
    
    doc.add_heading('5.2 ماذا يرفض الروبوت؟', level=2)
    p = doc.add_paragraph()
    run = p.add_run('يرفض الروبوت الأسئلة خارج النطاق التقني، مثل:')
    run.font.name = 'Cairo'
    
    non_technical = [
        'الأسئلة الشخصية أو الاجتماعية',
        'الاستفسارات الدينية أو السياسية',
        'المواضيع الطبية أو القانونية',
        'الأسئلة العامة غير المتعلقة بالتكنولوجيا'
    ]
    
    for item in non_technical:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('5.3 سياسة كلمة التنبيه', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'يجب أن تبدأ جميع الأسئلة بكلمة "روبوت" لكي يستجيب النظام. '
        'مثال: "روبوت ما هي بايثون؟" - سيتم الرد عليه. '
        'أما "ما هي بايثون؟" بدون كلمة الروبوت - سيتم تجاهله.'
    )
    run.font.name = 'Cairo'
    
    # Add callout box
    p = doc.add_paragraph()
    run = p.add_run('💡 ملاحظة: يمكن تعديل كلمة التنبيه في إعدادات النظام حسب الحاجة.')
    run.font.name = 'Cairo'
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    add_page_break(doc)
    
    # ============ 6. STAKEHOLDERS ============
    doc.add_heading('6. المستخدمون المستهدفون', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('يستهدف المشروع الفئات التالية:')
    run.font.name = 'Cairo'
    
    # Create table for stakeholders
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'الفئة المستهدفة'
    hdr_cells[1].text = 'الاحتياجات'
    hdr_cells[2].text = 'الفائدة المتوقعة'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    # Data rows
    stakeholders = [
        ('الطلاب', 'مساعدة في الواجبات والمشاريع', 'توفير الوقت وتعلم أسرع'),
        ('أعضاء الهيئة التدريسية', 'إجابات سريعة للاستفسارات التقنية', 'دعم في التدريس والبحث'),
        ('فريق الدعم الفني', 'تقليل حمول العمل المتكرر', 'التركيز على المشاكل المعقدة'),
        ('الباحثون', 'معلومات تقنية دقيقة', 'دعم في الأبحاث والتطوير'),
    ]
    
    for stakeholder, needs, benefits in stakeholders:
        row_cells = table.add_row().cells
        row_cells[0].text = stakeholder
        row_cells[1].text = needs
        row_cells[2].text = benefits
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 7. REQUIREMENTS ============
    doc.add_heading('7. المتطلبات', level=1)
    
    doc.add_heading('7.1 المتطلبات الوظيفية', level=2)
    
    functional_reqs = [
        'القدرة على قبول المدخلات النصية والصوتية',
        'التعرف على كلمة التنبيه "روبوت"',
        'معالجة الأسئلة التقنية باللغة العربية',
        'توليد إجابات دقيقة ومناسبة',
        'تحويل الإجابات النصية إلى صوت',
        'عرض سجل المحادثة السابقة',
        'رفض الأسئلة خارج النطاق التقني بلطف',
        'دعم اتصالات WebSocket للاتصال الفوري'
    ]
    
    for i, req in enumerate(functional_reqs, 1):
        p = doc.add_paragraph(f'FR-{i}: {req}')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('7.2 المتطلبات غير الوظيفية', level=2)
    
    # Create table for non-functional requirements
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'النوع'
    hdr_cells[1].text = 'المتطلب'
    hdr_cells[2].text = 'الوصف'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    non_functional = [
        ('الأداء', 'زمن استجابة < 3 ثوانٍ', 'ردود سريعة لتجربة مستخدم جيدة'),
        ('الخصوصية', 'عدم تخزين المحادثات', 'حماية بيانات المستخدمين'),
        ('سهولة الاستخدام', 'واجهة بسيطة', 'لا حاجة لتدريب مسبق'),
        ('التوافقية', 'Chrome/Edge', 'دعم المتصفحات الحديثة'),
        ('الأمان', 'API Key محمي', 'عدم暴露 المفاتيح في الواجهة'),
        ('قابلية التوسع', 'معمارية معيارية', 'سهولة إضافة ميزات جديدة'),
    ]
    
    for nfr_type, req, desc in non_functional:
        row_cells = table.add_row().cells
        row_cells[0].text = nfr_type
        row_cells[1].text = req
        row_cells[2].text = desc
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 8. TECHNOLOGIES AND TOOLS ============
    doc.add_heading('8. التقنيات والأدوات المستخدمة', level=1)
    
    # Create comprehensive technology table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'التقنية / الأداة'
    hdr_cells[1].text = 'الغرض'
    hdr_cells[2].text = 'الإصدار الموصى به'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    technologies = [
        ('JavaScript', 'لغة البرمجة الرئيسية', 'ES6+'),
        ('Node.js', 'بيئة التشغيل الخلفية', 'v18+'),
        ('Express', 'إطار عمل الويب', 'v4.18+'),
        ('ws', 'مكتبة WebSocket', 'v8+'),
        ('HTML5', 'بنية الواجهة الأمامية', 'HTML5'),
        ('CSS3', 'تصميم الواجهة', 'CSS3'),
        ('Web Speech API', 'التعرف على الصوت والتوليد الصوتي', 'Native Browser API'),
        ('Vertex AI / Gemini', 'نموذج الذكاء الاصطناعي', 'Latest'),
        ('dotenv', 'إدارة متغيرات البيئة', 'v16+'),
        ('Git', 'التحكم في الإصدارات', 'v2+'),
        ('VS Code', 'بيئة التطوير', 'Latest'),
    ]
    
    for tech, purpose, version in technologies:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose
        row_cells[2].text = version
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
    
    doc.add_paragraph().add_run().add_break()
    
    doc.add_heading('8.1 هيكل المشروع', level=2)
    p = doc.add_paragraph()
    run = p.add_run('يتكون المشروع من المجلدات والملفات التالية:')
    run.font.name = 'Cairo'
    
    structure = '''
    bot_it/
    ├── backend/
    │   ├── config/
    │   │   └── gemini.js       # إعدادات Vertex AI
    │   ├── server.js           # خادم WebSocket الرئيسي
    │   └── handlers/
    │       └── messageHandler.js  # معالجة الرسائل
    ├── frontend/
    │   ├── index.html          # الصفحة الرئيسية
    │   ├── styles.css          # التنسيقات
    │   └── app.js              # منطق التطبيق
    ├── .env                    # متغيرات البيئة
    ├── package.json            # تبعيات المشروع
    └── README.md               # التوثيق
    '''
    
    p = doc.add_paragraph(structure)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)
    
    add_page_break(doc)
    
    # ============ 9. ARCHITECTURE ============
    doc.add_heading('9. المعمارية', level=1)
    
    doc.add_heading('9.1 طبقات النظام', level=2)
    p = doc.add_paragraph()
    run = p.add_run('يتكون النظام من ثلاث طبقات رئيسية:')
    run.font.name = 'Cairo'
    
    layers = [
        ('طبقة العرض (Presentation Layer)', 'واجهة المستخدم في المتصفح'),
        ('طبقة المنطق (Business Logic Layer)', 'خادم Node.js ومعالجة الرسائل'),
        ('طبقة الخدمات (Services Layer)', 'Vertex AI API وWeb Speech API'),
    ]
    
    for layer, desc in layers:
        p = doc.add_paragraph()
        run = p.add_run(f'• {layer}: {desc}')
        run.font.name = 'Cairo'
    
    doc.add_heading('9.2 تدفق البيانات', level=2)
    p = doc.add_paragraph()
    run = p.add_run('يتم تدفق البيانات عبر الخطوات التالية:')
    run.font.name = 'Cairo'
    
    flow_steps = [
        '1. المستخدم يضغط على زر التحدث وينطق السؤال',
        '2. Web Speech API (SpeechRecognition) يحول الصوت إلى نص',
        '3. النص يُرسل عبر WebSocket إلى الخادم',
        '4. الخادم يتحقق من وجود كلمة التنبيه "روبوت"',
        '5. إذا وُجدت الكلمة، يُرسل السؤال إلى Vertex AI/Gemini',
        '6. النموذج يولد إجابة نصية مناسبة',
        '7. الإجابة تُعاد عبر WebSocket إلى الواجهة',
        '8. Web Speech API (SpeechSynthesis) يحول النص إلى صوت',
        '9. يتم عرض الإجابة نصياً وتشغيلها صوتياً'
    ]
    
    for step in flow_steps:
        p = doc.add_paragraph(step)
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('9.3 مخطط المعمارية', level=2)
    p = doc.add_paragraph()
    run = p.add_run('''
    ┌─────────────────────────────────────────────────────────────┐
    │                    طبقة العرض (Frontend)                    │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
    │  │   HTML/CSS   │  │   Web Speech │  │  WebSocket   │      │
    │  │   Interface  │  │     API      │  │   Client     │      │
    │  └──────────────┘  └──────────────┘  └──────────────┘      │
    └──────────────────────────────┬──────────────────────────────┘
                                 │ WebSocket
                                 ▼
    ┌─────────────────────────────────────────────────────────────┐
    │              طبقة المنطق (Backend - Node.js)                │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐      │
    │  │  Express     │  │   Wake Word  │  │   Message    │      │
    │  │   Server     │  │    Check     │  │   Handler    │      │
    │  └──────────────┘  └──────────────┘  └──────────────┘      │
    └──────────────────────────────┬──────────────────────────────┘
                                 │ HTTPS
                                 ▼
    ┌─────────────────────────────────────────────────────────────┐
    │              طبقة الخدمات (AI Services)                     │
    │  ┌──────────────────────────────────────────────────────┐  │
    │  │            Vertex AI / Gemini API                     │  │
    │  │           (Text Generation Model)                     │  │
    │  └──────────────────────────────────────────────────────┘  │
    └─────────────────────────────────────────────────────────────┘
    ''')
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    add_page_break(doc)
    
    # ============ 10. UI/UX DESIGN ============
    doc.add_heading('10. تصميم واجهة المستخدم', level=1)
    
    doc.add_heading('10.1 المبادئ الأساسية', level=2)
    
    principles = [
        'دعم RTL (من اليمين إلى اليسار) للعربية',
        'تصميم بسيط وواضح',
        'ألوان هادئة ومريحة للعين',
        'تجربة مستخدم سلسة',
        'مؤشرات حالة واضحة'
    ]
    
    for principle in principles:
        p = doc.add_paragraph(f'• {principle}')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('10.2 مكونات الواجهة', level=2)
    
    # Create table for UI components
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'المكون'
    hdr_cells[1].text = 'الوصف'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    ui_components = [
        ('زر Push-to-Talk', 'زر كبير واضح للتحدث'),
        ('سجل المحادثة', 'عرض الأسئلة والإجابات السابقة'),
        ('مؤشر الحالة', 'يعرض حالة الاستماع/التفكير/التحدث'),
        ('إعدادات الصوت', 'التحكم في مستوى الصوت وسرعة الكلام'),
        ('معلومات النظام', 'عرض حالة الاتصال'),
    ]
    
    for component, desc in ui_components:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = desc
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
    
    doc.add_heading('10.3 مخطط الواجهة', level=2)
    p = doc.add_paragraph()
    run = p.add_run('''
    ┌────────────────────────────────────────────────────────────┐
    │                    Bot_IT - روبوت الجامعة التقني           │
    ├────────────────────────────────────────────────────────────┤
    │  ┌──────────────────────────────────────────────────────┐ │
    │  │  سجل المحادثة:                                       │ │
    │  │  أنت: روبوت ما هي بايثون؟                          │ │
    │  │  الروبوت: بايثون هي لغة برمجة...                   │ │
    │  │                                                      │ │
    │  └──────────────────────────────────────────────────────┘ │
    │                                                            │
    │                   ┌──────────────┐                         │
    │                   │ 🎤 اضغط للتحدث │                         │
    │                   └──────────────┘                         │
    │                                                            │
    │  الحالة: جاهز | الاتصال: متصل                             │
    └────────────────────────────────────────────────────────────┘
    ''')
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    add_page_break(doc)
    
    # ============ 11. SECURITY AND PRIVACY ============
    doc.add_heading('11. الأمان والخصوصية', level=1)
    
    doc.add_heading('11.1 حماية API Keys', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'يتم تخزين مفاتيح API في ملف .env على الخادم فقط، ولا يتم إرسالها أبداً '
        'إلى الواجهة الأمامية. هذا يمنع الوصول غير المصرح به إلى حساب Vertex AI.'
    )
    run.font.name = 'Cairo'
    
    # Add warning callout
    p = doc.add_paragraph()
    run = p.add_run('⚠️ تحذير: لا تضع أبداً API Keys في الكود الأمامي أو في مستودع Git عام!')
    run.font.name = 'Cairo'
    run.font.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    
    doc.add_heading('11.2 خصوصية البيانات', level=2)
    
    privacy_points = [
        'لا يتم تخزين محادثات المستخدمين بشكل دائم',
        'البيانات تُرسل مباشرة إلى Vertex AI عبر HTTPS',
        'Web Speech API قد يستخدم خدمات سحابية للمتصفح',
        'يجب استخدام HTTPS عند النشر لحماية البيانات أثناء النقل'
    ]
    
    for point in privacy_points:
        p = doc.add_paragraph(f'• {point}')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('11.3 توصيات الأمان', level=2)
    
    recommendations = [
        'تقييد مفاتيح API بـ IP addresses محددة',
        'استخدام Rate Limiting لمنع الاستهلاك المفرط',
        'تطبيق CORS للتحكم في النطاقات المسموح بها',
        'تشغيل الخادم خلف جدار حماية',
        'تحديث التبعيات بانتظام'
    ]
    
    for rec in recommendations:
        p = doc.add_paragraph(f'• {rec}')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 12. TESTING PLAN ============
    doc.add_heading('12. خطة الاختبار', level=1)
    
    doc.add_heading('12.1 حالات الاختبار', level=2)
    
    # Create comprehensive test cases table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'رقم الاختبار'
    hdr_cells[1].text = 'الحالة'
    hdr_cells[2].text = 'المدخلات'
    hdr_cells[3].text = 'النتيجة المتوقعة'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    test_cases = [
        ('TC-01', 'مع كلمة التنبيه', 'روبوت ما هي بايثون؟', 'إجابة عن بايثون'),
        ('TC-02', 'بدون كلمة التنبيه', 'ما هي بايثون؟', 'لا رد أو رسالة تذكير'),
        ('TC-03', 'سؤال تقني صحيح', 'روبوت كيف أعمل دالة في JavaScript؟', 'إجابة تقنية دقيقة'),
        ('TC-04', 'سؤال خارج النطاق', 'روبوت كيف حالك؟', 'رفض لطيف وتوجيه للأسئلة التقنية'),
        ('TC-05', 'مدخل صوتي', 'نطق "روبوت ما هو HTTP؟"', 'تحويل صحيح وإجابة'),
        ('TC-06', 'اتصال WebSocket', 'فتح الصفحة', 'اتصال ناجح'),
        ('TC-07', 'إخراج صوتي', 'أي سؤال صحيح', 'تشغيل الإجابة صوتياً'),
        ('TC-08', 'سجل المحادثة', 'عدة أسئلة متتالية', 'عرض السجل كاملاً'),
    ]
    
    for test_id, case, input_val, expected in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = test_id
        row_cells[1].text = case
        row_cells[2].text = input_val
        row_cells[3].text = expected
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
                    run.font.size = Pt(9)
    
    doc.add_heading('12.2 اختبار المتصفحات', level=2)
    p = doc.add_paragraph()
    run = p.add_run('تم اختبار النظام على المتصفحات التالية:')
    run.font.name = 'Cairo'
    
    browsers = [
        ('Google Chrome', '✅ مدعوم بالكامل'),
        ('Microsoft Edge', '✅ مدعوم بالكامل'),
        ('Mozilla Firefox', '⚠️ دعم جزئي للـ Web Speech API'),
        ('Safari', '❌ غير مدعوم حالياً'),
    ]
    
    for browser, status in browsers:
        p = doc.add_paragraph(f'• {browser}: {status}')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 13. EXPECTED RESULTS ============
    doc.add_heading('13. النتائج المتوقعة ومؤشرات النجاح', level=1)
    
    doc.add_heading('13.1 النتائج المتوقعة', level=2)
    
    expected_results = [
        'نظام روبوت محادثة فعال يعمل على localhost',
        'استجابة سريعة للاستفسارات التقنية (أقل من 3 ثوانٍ)',
        'دقة عالية في الإجابات (تتجاوز 85%)',
        'واجهة مستخدم سهلة وجذابة',
        'دعم كامل للغة العربية (نصاً وصوتاً)'
    ]
    
    for result in expected_results:
        p = doc.add_paragraph(f'✓ {result}')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('13.2 مؤشرات النجاح', level=2)
    
    # Create table for success metrics
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'المؤشر'
    hdr_cells[1].text = 'الهدف'
    hdr_cells[2].text = 'كيفية القياس'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    metrics = [
        ('دقة الإجابات', '> 85%', 'اختبار يدوي لعينة من الأسئلة'),
        ('زمن الاستجابة', '< 3 ثوانٍ', 'قياس الوقت من السؤال للإجابة'),
        ('رضا المستخدمين', '> 4/5', 'استبيان بعد الاستخدام'),
        ('نسبة النجاح', '> 90%', 'نسبة الأسئلة المجاب عليها بنجاح'),
    ]
    
    for metric, target, measurement in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = target
        row_cells[2].text = measurement
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 14. CHALLENGES AND SOLUTIONS ============
    doc.add_heading('14. التحديات والمشاكل التي واجهت المشروع والحلول', level=1)
    
    # Create challenges table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'التحدي'
    hdr_cells[1].text = 'التأثير'
    hdr_cells[2].text = 'الحل المطبق'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
    
    challenges = [
        (
            'دعم TTS العربية',
            'عدم توفر أصوات عربية جاهزة',
            'استخدام Google Chrome وتثبيت حزم اللغة العربية'
        ),
        (
            'دقة التعرف على الصوت',
            'أخطاء في تحويل الكلام العربي',
            'استخدام Web Speech API مع Chrome وتحسين جودة الميكروفون'
        ),
        (
            'حدود نطاق الأسئلة',
            'صعوبة تحديد الأسئلة التقنية',
            'تطبيق prompt engineering وسياسة واضحة للنطاق'
        ),
        (
            'تأخير الاستجابة',
            'بطء في الاتصال بـ Vertex AI',
            'تحسين كود الخادم واستخدام WebSocket بدلاً من HTTP'
        ),
        (
            'أمان API Keys',
            'خطر التسريب في الكود',
            'استخدام ملف .env وعدم إرسال المفاتيح للواجهة'
        ),
    ]
    
    for challenge, impact, solution in challenges:
        row_cells = table.add_row().cells
        row_cells[0].text = challenge
        row_cells[1].text = impact
        row_cells[2].text = solution
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
                    run.font.size = Pt(10)
    
    add_page_break(doc)
    
    # ============ 15. FUTURE IMPROVEMENTS ============
    doc.add_heading('15. التحسينات المستقبلية', level=1)
    
    improvements = [
        ('النشر على سيرفر سحابي', 'إتاحة الوصول من أي مكان'),
        ('لوحة تحكم إدارية', 'لإحصائيات الاستخدام وإدارة المحتوى'),
        ('قاعدة معرفة مخصصة', 'لإضافة معلومات خاصة بالجامعة'),
        ('ربط RAG', 'لتحسين دقة الإجابات باستخدام Retrieval-Augmented Generation'),
        ('دعم تعدد اللغات', 'إضافة الإنجليزية والفرنسية'),
        ('تطبيق موبايل', 'لأجهزة Android وiOS'),
        ('إعدادات شخصية', 'للتحكم في سرعة الصوت ونبرة الكلام'),
        ('نظام تسجيل دخول', 'لتخصيص التجربة وحفظ السجل'),
        ('تكامل مع أنظمة الجامعة', 'مثل نظام إدارة التعلم LMS'),
        ('وضع متعدد المستخدمين', 'للمحادثات الجماعية'),
    ]
    
    for i, (improvement, desc) in enumerate(improvements, 1):
        p = doc.add_paragraph(f'{i}. {improvement}: {desc}')
        p.runs[0].font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 16. CONCLUSION ============
    doc.add_heading('16. خاتمة', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'يمثل مشروع Bot_IT خطوة مهمة نحو تحسين الدعم التقني في البيئة الجامعية. '
        'من خلال الجمع بين تقنيات الذكاء الاصطناعي الحديثة وواجهة مستخدم بسيطة، '
        'يوفر المشروع حلاً فعالاً ومبتكراً للاستفسارات التقنية.'
    )
    run.font.name = 'Cairo'
    
    p = doc.add_paragraph()
    run = p.add_run(
        'لقد تم تصميم النظام مع مراعاة مبادئ الأمان والخصوصية، مع قابلية التوسع '
        'والتطوير مستقبلاً. نتوقع أن يساهم المشروع في تحسين تجربة التعلم '
        'وتقليل العبء على أعضاء الهيئة التدريسية وفريق الدعم الفني.'
    )
    run.font.name = 'Cairo'
    
    p = doc.add_paragraph()
    run = p.add_run(
        'نأمل أن يكون هذا المشروع نقطة انطلاق لمزيد من الابتكارات في مجال '
        'التعليم الذكي والدعم الأكاديمي الآلي.'
    )
    run.font.name = 'Cairo'
    
    add_page_break(doc)
    
    # ============ 17. APPENDIX ============
    doc.add_heading('17. ملاحق', level=1)
    
    doc.add_heading('17.1 هيكل الملفات الكامل', level=2)
    p = doc.add_paragraph()
    run = p.add_run('''
    bot_it/
    │
    ├── backend/
    │   ├── config/
    │   │   └── gemini.js           # إعدادات Vertex AI
    │   ├── server.js               # خادم WebSocket الرئيسي
    │   └── handlers/
    │       └── messageHandler.js   # معالجة الرسائل
    │
    ├── frontend/
    │   ├── index.html              # الصفحة الرئيسية
    │   ├── styles.css              # التنسيقات
    │   └── app.js                  # منطق التطبيق
    │
    ├── .env                        # متغيرات البيئة (API Keys)
    ├── .gitignore                  # ملفات Git المتجاهلة
    ├── package.json                # تبعيات المشروع
    ├── README.md                   # التوثيق
    └── generate_report.py          # هذا الملف
    ''')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('17.2 أوامر التشغيل', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('لتشغيل المشروع، اتبع الخطوات التالية:')
    run.font.name = 'Cairo'
    run.font.bold = True
    
    commands = '''
    1. تثبيت التبعيات:
       npm install
    
    2. إعداد ملف .env:
       GEMINI_API_KEY=your_api_key_here
       PORT=3000
    
    3. تشغيل الخادم:
       npm start
    
    4. فتح المتصفح:
       افتح http://localhost:3000 في Google Chrome
    '''
    
    p = doc.add_paragraph(commands)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_heading('17.3 المتطلبات الأساسية', level=2)
    
    prerequisites = [
        'Node.js (v18 أو أحدث)',
        'npm (v9 أو أحدث)',
        'Google Chrome أو Microsoft Edge',
        'مفتاح API من Vertex AI/Gemini',
        'اتصال إنترنت نشط'
    ]
    
    for prereq in prerequisites:
        p = doc.add_paragraph(f'• {prereq}')
        p.runs[0].font.name = 'Cairo'
    
    doc.add_heading('17.4 ملاحظات مهمة', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('⚠️ تأكد من:')
    run.font.name = 'Cairo'
    run.font.bold = True
    
    notes = [
        'عدم مشاركة ملف .env مع أي شخص',
        'استخدام HTTPS عند النشر للإنتاج',
        'مراجعة أسعار Vertex AI قبل الاستخدام المكثف',
        'اختبار النظام جيداً قبل الإطلاق'
    ]
    
    for note in notes:
        p = doc.add_paragraph(f'• {note}')
        p.runs[0].font.name = 'Cairo'
    
    # Add page numbers to all sections
    for section in doc.sections:
        add_page_number(section)
    
    # Save the document
    output_path = 'Bot_IT_Project_Report.docx'
    doc.save(output_path)
    
    print("Report created successfully: " + output_path)
    print("Path: " + os.path.abspath(output_path))
    print("Expected pages: 16-18")
    print("Language: Arabic (RTL)")
    print("Format: Microsoft Word (.docx)")
    
    return output_path

# Import WD_ORIENTATION at the top level
from docx.enum.section import WD_ORIENTATION

if __name__ == '__main__':
    try:
        create_bot_it_report()
    except Exception as e:
        print("Error creating report: " + str(e))
        import traceback
        traceback.print_exc()
