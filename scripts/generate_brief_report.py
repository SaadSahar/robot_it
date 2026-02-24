#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Bot_IT Brief Report Generator
Generates a concise Word document for the university robot project
Based on doc_pro2.md specifications
"""

import sys
import io
# Set UTF-8 encoding for output
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL (Right-to-Left) for Arabic text"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        run.font.rtl = True

def create_bot_it_brief_report():
    """Create the brief Bot_IT project report (6 sections, 4-6 pages)"""
    
    # Create a new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "تقرير مشروع روبوت الجامعة التقني Bot_IT"
    doc.core_properties.author = "فريق المشروع"
    doc.core_properties.subject = "Bot_IT University Robot Project"
    
    # Set up styles
    styles = doc.styles
    
    # Normal style - Arabic font
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Cairo'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Heading 1 style
    heading1 = styles['Heading 1']
    heading1.font.name = 'Cairo'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Heading 2 style
    heading2 = styles['Heading 2']
    heading2.font.name = 'Cairo'
    heading2.font.size = Pt(16)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 76, 153)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ============ COVER PAGE ============
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('اسم الجامعة: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Project title
    doc.add_paragraph().add_run().add_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('تقرير مشروع')
    run.font.name = 'Cairo'
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('روبوت الجامعة التقني Bot_IT')
    run.font.name = 'Cairo'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 76, 153)
    
    # Team members
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('أعضاء الفريق:')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    run.font.bold = True
    
    for i in range(1, 5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{i}. ____________________')
        run.font.name = 'Cairo'
        run.font.size = Pt(12)
    
    # Supervisor
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('المشرف: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    run.font.bold = True
    
    # Date
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('التاريخ: ____________________')
    run.font.name = 'Cairo'
    run.font.size = Pt(13)
    
    # Add page break after cover
    doc.add_page_break()
    
    # ============ SECTION 1: فكرة المشروع ============
    doc.add_heading('1. فكرة المشروع', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'يُعد Bot_IT روبوتاً تفاعلياً صوتياً موجوداً في الجامعة، مصمم للإجابة على '
        'الأسئلة التقنية والتكنولوجية مثل البرمجة والشبكات والذكاء الاصطناعي. '
        'يساعد الروبوت الطلاب في الحصول على إجابات سريعة ودقيقة، مع رفض الأسئلة '
        'غير التقنية بلطف. يتم تفعيل الروبوت بكلمة "روبوت" قبل السؤال، مثل: '
        '"روبوت ما هي لغة بايثون؟".'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    doc.add_paragraph().add_run().add_break()
    
    # ============ SECTION 2: أهداف المشروع ============
    doc.add_heading('2. أهداف المشروع', level=1)
    
    objectives = [
        'توفير مساعد تقني ذكي للطلاب على مدار الساعة',
        'تقديم إجابات فورية على الاستفسارات التقنية',
        'تسهيل الوصول للمعلومات التقنية الموثوقة',
        'تقليل الضغط على المختبرات والمشرفين',
        'تطبيق تقنيات الذكاء الاصطناعي عملياً في البيئة الجامعية',
        'تحسين تجربة التعلم الذاتي للطلاب'
    ]
    
    for objective in objectives:
        p = doc.add_paragraph(objective, style='List Bullet')
        p.runs[0].font.name = 'Cairo'
        p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph().add_run().add_break()
    
    # ============ SECTION 3: تعريف بسيط عن الذكاء الاصطناعي ============
    doc.add_heading('3. تعريف بسيط عن الذكاء الاصطناعي', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'الذكاء الاصطناعي (Artificial Intelligence) هو فرع من علوم الحاسوب يهتم '
        'بتطوير أنظمة قادرة على محاكاة الذكاء البشري. تُمكّن هذه التقنية الآلات '
        'من التفكير والتعلم واتخاذ القرارات بناءً على البيانات. من أمثلة الذكاء '
        'الاصطناعي في حياتنا اليومية: المساعدات الصوتية مثل Siri وGoogle Assistant، '
        'أنظمة التوصيات على Netflix وYouTube، والمركبات ذاتية القيادة.'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    doc.add_paragraph().add_run().add_break()
    
    # ============ SECTION 4: شرح مختصر عن نماذج اللغة الكبيرة (LLM) ============
    doc.add_heading('4. شرح مختصر عن نماذج اللغة الكبيرة (LLM)', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'نماذج اللغة الكبيرة (Large Language Models) هي أنظمة ذكاء اصطناعي متقدمة '
        'تم تدريبها على كميات ضخمة من النصوص لفهم اللغة البشرية وتوليد ردود طبيعية. '
        'تستخدم هذه النماذج تقنيات التعلم العميق لفهم سياق الأسئلة وتقديم إجابات '
        'منطقية ومترابطة. يستخدم روبوت Bot_IT نموذج Gemini من Google كمحرك للذكاء '
        'الاصطناعي لفهم أسئلة الطلاب وتوليد الإجابات المناسبة.'
    )
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    doc.add_paragraph().add_run().add_break()
    
    # ============ SECTION 5: التقنيات المستخدمة ============
    doc.add_heading('5. التقنيات المستخدمة', level=1)
    
    # Create table for technologies
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'التقنية'
    hdr_cells[1].text = 'الوصف'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Cairo'
                run.font.bold = True
                run.font.size = Pt(12)
    
    # Data rows
    technologies = [
        ('JavaScript', 'لغة البرمجة الأساسية للمشروع'),
        ('Node.js', 'بيئة تشغيل الكود على مركز المعالجة'),
        ('Gemini AI', 'نموذج الذكاء الاصطناعي لفهم الأسئلة وتوليد الإجابات'),
        ('Web Speech API', 'تحويل الصوت إلى نص والنص إلى صوت'),
        ('WebSocket', 'بروتوكول الاتصال السريع بين مكونات الروبوت')
    ]
    
    for tech, description in technologies:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = description
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cairo'
                    run.font.size = Pt(11)
    
    doc.add_paragraph().add_run().add_break()
    
    # ============ SECTION 6: آلية عمل الروبوت ============
    doc.add_heading('6. آلية عمل الروبوت', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('تعتمد آلية عمل الروبوت على الخطوات التالية:')
    run.font.name = 'Cairo'
    run.font.size = Pt(12)
    
    mechanism_steps = [
        'المستخدم يقول سؤاله للروبوت (عبر المايكروفون)',
        'الروبوت يحوّل الصوت إلى نص',
        'يتحقق النظام من وجود كلمة التنبيه "روبوت"',
        'يُرسل السؤال إلى نموذج الذكاء الاصطناعي (Gemini)',
        'يستقبل النظام الإجابة ويحوّلها إلى صوت',
        'الروبوت ينطق الإجابة للمستخدم (عبر السماعة)'
    ]
    
    for i, step in enumerate(mechanism_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.runs[0].font.name = 'Cairo'
        p.runs[0].font.size = Pt(12)
    
    # Save the document
    output_path = 'Bot_IT_Report_Brief.docx'
    doc.save(output_path)
    print(f"✓ Report generated successfully: {output_path}")
    print(f"✓ Total sections: 6")
    print(f"✓ Format: RTL (Arabic)")
    print(f"✓ Focus: Physical robot (not website)")
    
    return output_path

if __name__ == "__main__":
    create_bot_it_brief_report()
